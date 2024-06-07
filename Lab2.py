from docx import Document
from docx.shared import Pt

def change_font_size_at_indices(docx_file, target_indices, new_size, new_docx_file):
    doc = Document(docx_file)
    new_doc = Document()  # Создаем новый объект документа

    char_index = 0
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()  # Добавляем новый абзац в новый документ
        for run in paragraph.runs:
            for i, char in enumerate(run.text):
                if char != "\n":
                    char_index += 1
                    if char_index - 1 in target_indices:
                        # Создаем новый фрагмент текста для символа с новым размером шрифта
                        new_run = new_paragraph.add_run(char)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = Pt(new_size)
                        target_indices.remove(char_index - 1)  # Удаляем индекс из списка целевых индексов
                    else:
                        new_paragraph.add_run(char)  # Добавляем символ в новый абзац как есть
                else:
                    new_paragraph.add_run("\n")  # Добавляем символ новой строки

    new_doc.save(new_docx_file)  # Сохраняем новый документ

# Пример использования:
change_font_size_at_indices("3.docx", [0, 3, 5, 6, 7, 8, 9, 10, 14, 16, 18, 20, 21, 22, 26, 32, 34, 36, 37, 38, 39, 40, 42, 44, 45, 46, 48, 49, 50, 54, 56, 57, 58, 64, 65, 66, 70, 71, 72, 74, 77, 80, 82, 84, 88, 90, 92, 94, 95, 96, 97, 98, 103, 104, 105, 106, 108, 109, 110, 111, 114, 116, 117, 122, 128, 129, 130, 134, 136, 138, 140, 141, 142, 146, 152, 154, 156, 162, 168, 170, 172, 173, 174, 175, 176, 178, 180, 181, 182, 184, 186, 189, 191, 192, 194, 196, 198, 199, 202, 204, 205, 206],11.5, "modified_doc.docx")
